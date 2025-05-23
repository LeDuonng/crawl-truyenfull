import requests
from bs4 import BeautifulSoup
from docx import Document
import re
import os
import concurrent.futures
import time
from pathlib import Path
import logging
from datetime import datetime

def setup_logging(story_name):
    """Thiết lập logging cho quá trình crawl"""
    log_dir = Path("logs")
    log_dir.mkdir(exist_ok=True)
    
    log_file = log_dir / f"{story_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
    
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(log_file, encoding='utf-8'),
            logging.StreamHandler()
        ]
    )
    return logging.getLogger(__name__)

def get_last_chapter(base_url, logger):
    """Tìm chương cuối cùng của truyện"""
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
        }
        response = requests.get(base_url, headers=headers, verify=False, timeout=30)
        
        if response.status_code == 200:
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Tìm link cuối theo nhiều cách
            last_page_links = []
            
            # Cách 1: Tìm theo string chứa "Cuối"
            try:
                last_page_link = soup.find('a', string=lambda t: t and 'Cuối' in t)
                if last_page_link:
                    last_page_links.append(last_page_link)
            except Exception as e:
                logger.warning(f"Lỗi khi tìm link 'Cuối' theo cách 1: {e}")
            
            # Cách 2: Tìm theo title chứa "Cuối"
            try:
                last_page_link = soup.find('a', title=lambda t: t and 'Cuối' in t)
                if last_page_link:
                    last_page_links.append(last_page_link)
            except Exception as e:
                logger.warning(f"Lỗi khi tìm link 'Cuối' theo cách 2: {e}")
            
            # Cách 3: Tìm theo class "last"
            try:
                last_page_link = soup.find('a', class_='last')
                if last_page_link:
                    last_page_links.append(last_page_link)
            except Exception as e:
                logger.warning(f"Lỗi khi tìm link 'Cuối' theo cách 3: {e}")
                
            # Cách 4: Tìm theo span.arrow
            try:
                arrow_spans = soup.find_all('span', class_='arrow')
                for span in arrow_spans:
                    if span.get_text() == '»' and span.parent.name == 'a':
                        last_page_links.append(span.parent)
                        break
            except Exception as e:
                logger.warning(f"Lỗi khi tìm link 'Cuối' theo cách 4: {e}")
                
            # Cách 5: Tìm link phân trang có số trang lớn nhất
            try:
                # Tìm tất cả link phân trang theo pattern trang-X
                pagination_links = soup.find_all('a', href=re.compile(r'/trang-\d+/'))
                if pagination_links:
                    # Lấy số trang từ mỗi link
                    max_page = 0
                    max_page_link = None
                    for link in pagination_links:
                        href = link.get('href', '')
                        page_match = re.search(r'/trang-(\d+)/', href)
                        if page_match:
                            page_num = int(page_match.group(1))
                            if page_num > max_page:
                                max_page = page_num
                                max_page_link = link
                    
                    if max_page_link:
                        logger.info(f"Tìm thấy trang cuối từ phân trang: trang-{max_page}")
                        last_page_links.append(max_page_link)
            except Exception as e:
                logger.warning(f"Lỗi khi tìm link phân trang: {e}")
            
            # Xử lý các link tìm được
            for last_page_link in last_page_links:
                last_page_url = last_page_link.get('href')
                if last_page_url:
                    logger.info(f"Tìm thấy link trang cuối: {last_page_url}")
                    # Truy cập trang cuối
                    last_page_response = requests.get(last_page_url, headers=headers, verify=False, timeout=30)
                    if last_page_response.status_code == 200:
                        last_page_soup = BeautifulSoup(last_page_response.content, 'html.parser')
                        
                        # Tìm chương cuối trong trang cuối
                        chapter_links = last_page_soup.find_all('a', href=re.compile(r'chuong-\d+'))
                        if chapter_links:
                            logger.info(f"Số link chương tìm thấy trong trang cuối: {len(chapter_links)}")
                            last_chapter_url = chapter_links[-1].get('href')
                            logger.info(f"URL chương cuối: {last_chapter_url}")
                            chapter_match = re.search(r'chuong-(\d+)', last_chapter_url)
                            if chapter_match:
                                last_chapter_number = int(chapter_match.group(1))
                                logger.info(f"Tìm thấy chương cuối từ trang cuối: {last_chapter_number}")
                                return last_chapter_number
                    else:
                        logger.warning(f"Không thể truy cập trang cuối: {last_page_url}, mã lỗi: {last_page_response.status_code}")
            
            # Nếu không tìm thấy link "Cuối", tìm trong danh sách chương
            chapter_list = soup.find_all('a', href=re.compile(r'chuong-\d+'))
            if chapter_list:
                last_chapter_url = chapter_list[-1].get('href')
                chapter_match = re.search(r'chuong-(\d+)', last_chapter_url)
                if chapter_match:
                    last_chapter_number = int(chapter_match.group(1))
                    logger.info(f"Tìm thấy chương cuối từ danh sách: {last_chapter_number}")
                    return last_chapter_number
        
        logger.error(f"Không thể truy cập trang chủ: {base_url}")
        return None
        
    except Exception as e:
        logger.error(f"Lỗi khi tìm chương cuối: {str(e)}")
        return None

def get_chapter_content(url, max_retries=3):
    for attempt in range(max_retries):
        try:
            headers = {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
            }
            response = requests.get(url, headers=headers, verify=False, timeout=30)
            
            if response.status_code == 200:
                soup = BeautifulSoup(response.content, 'html.parser')
                chapter_div = soup.find('div', class_='chapter-c')
                
                if chapter_div:
                    for ad_div in chapter_div.find_all('div', class_=lambda x: x and 'ads' in x):
                        ad_div.extract()
                    
                    hr_end = chapter_div.find('hr', class_='chapter-end')
                    if hr_end:
                        content_elements = []
                        for sibling in hr_end.previous_siblings:
                            if sibling.name == 'div' and any('ads' in c for c in sibling.get('class', [])):
                                continue
                            content_elements.append(sibling)
                        content_elements.reverse()
                        
                        temp_html = ''.join(str(elem) for elem in content_elements)
                        temp_soup = BeautifulSoup(temp_html, 'html.parser')
                        text = temp_soup.get_text(separator='\n').strip()
                    else:
                        text = chapter_div.get_text(separator='\n').strip()
                    
                    # Kiểm tra nếu text rỗng hoặc chỉ chứa khoảng trắng
                    if not text or text.isspace():
                        print(f"Chương trống: {url}")
                        return "Chương trống"
                        
                    return text
                else:
                    print(f"Không tìm thấy div chapter-c: {url}")
                    return "Không tìm thấy nội dung"
            elif response.status_code == 503:
                if attempt < max_retries - 1:
                    wait_time = (attempt + 1) * 5
                    print(f"Lỗi 503, thử lại sau {wait_time} giây...")
                    time.sleep(wait_time)
                    continue
                else:
                    print(f"Không thể truy cập {url} sau {max_retries} lần thử")
                    return "Lỗi 503"
            else:
                print(f"Không thể truy cập {url}, mã lỗi: {response.status_code}")
                return f"Lỗi {response.status_code}"
        except Exception as e:
            if attempt < max_retries - 1:
                wait_time = (attempt + 1) * 5
                print(f"Lỗi khi lấy dữ liệu từ {url}: {e}, thử lại sau {wait_time} giây...")
                time.sleep(wait_time)
                continue
            else:
                print(f"Lỗi khi lấy dữ liệu từ {url}: {e}")
                return f"Lỗi: {str(e)}"
    return "Lỗi không xác định"

def detect_ads(text):
    """Nhận biết quảng cáo dựa trên các đặc điểm"""
    # Danh sách từ khóa quảng cáo
    ad_keywords = [
        'theo dõi', 'subscribe', 'follow', 'like', 'bình luận', 'comment',
        'review', 'đánh giá', 'kênh', 'channel', 'team', 'nhóm', 'group',
        'facebook', 'fb', 'fanpage', 'website', 'web', 'link'
    ]
    
    # Danh sách emoji phổ biến trong quảng cáo
    ad_emojis = ['📍', '🌺', '❤️', '💖', '💝', '💕', '💗', '💓', '💞', '💟', '💜', '💙', '💚', '💛', '🧡', '🤍', '🤎', '🖤', '💯', '⭐', '🌟', '✨']
    
    # Kiểm tra có emoji ở đầu và cuối không
    has_emoji_ends = any(text.startswith(emoji) and text.endswith(emoji) for emoji in ad_emojis)
    
    # Kiểm tra có từ khóa quảng cáo không
    has_ad_keywords = any(keyword in text.lower() for keyword in ad_keywords)
    
    # Kiểm tra có link không
    has_links = bool(re.search(r'https?://\S+|www\.\S+', text))
    
    # Kiểm tra có yêu cầu tương tác không
    has_interaction_requests = bool(re.search(r'(like|theo dõi|subscribe|bình luận|comment|review|đánh giá)', text.lower()))
    
    # Nếu có ít nhất 2 trong 4 đặc điểm trên, coi là quảng cáo
    return sum([has_emoji_ends, has_ad_keywords, has_links, has_interaction_requests]) >= 2

def clean_text(text):
    # Loại bỏ các dòng trống thừa và chuẩn hóa khoảng trắng
    text = re.sub(r'\n{2,}', '\n\n', text)  # Giữ tối đa 2 dòng trống giữa các đoạn
    text = re.sub(r'\s+', ' ', text)       # Chuẩn hóa khoảng trắng
    
    # Tách text thành các đoạn
    paragraphs = text.split('\n\n')
    
    # Lọc bỏ các đoạn quảng cáo
    cleaned_paragraphs = []
    for para in paragraphs:
        if not detect_ads(para.strip()):
            cleaned_paragraphs.append(para)
    
    # Ghép lại các đoạn
    text = '\n\n'.join(cleaned_paragraphs)
    
    # Loại bỏ các pattern quảng cáo cụ thể đã biết
    text = re.sub(r'\[Truyện được đăng tải duy nhất tại monkeydtruyen\.com - https://monkeydtruyen\.com/.*?\]', '', text)
    text = re.sub(r'📍 Nếu thấy hay đừng ngại cho bọn mình một lượt theo dõi nhé! 📍.*?Cá Chép Ngắm Mưa • 鯉魚望雨.*?hấp dẫn!', '', text, flags=re.DOTALL)
    text = re.sub(r'🌺 Hi, Chào mừng bạn ghé kênh của team Nhân Trí.*?Cảm ơn bạn 🌺', '', text, flags=re.DOTALL)
    
    return text.strip()

def fetch_chapter(chapter_number, base_url):
    # Thử URL chuẩn trước
    chapter_url = f"{base_url}chuong-{chapter_number}/"
    print(f"Đang lấy chương {chapter_number}: {chapter_url}")
    content = get_chapter_content(chapter_url)
    
    # Nếu không tìm thấy, thử các URL đặc biệt
    if content in ["Không tìm thấy nội dung", "Lỗi 404"]:
        # Thử URL với format chuong-XXX-YYY
        special_url = f"{base_url}chuong-{chapter_number}-{chapter_number+1}/"
        print(f"Thử URL đặc biệt: {special_url}")
        content = get_chapter_content(special_url)
        
        if content in ["Không tìm thấy nội dung", "Lỗi 404"]:
            # Thử URL với format chuong-XXX/YYY
            special_url = f"{base_url}chuong-{chapter_number}/{chapter_number+1}/"
            print(f"Thử URL đặc biệt: {special_url}")
            content = get_chapter_content(special_url)
    
    return (chapter_number, content)

def crawl_story(base_url, batch_size=10, max_workers=5):
    # Lấy tên truyện từ URL
    story_name = base_url.split('/')[-2].replace('-', '_')
    
    # Thiết lập logging
    logger = setup_logging(story_name)
    logger.info(f"Bắt đầu crawl truyện: {base_url}")
    
    # Tìm chương cuối
    last_chapter = get_last_chapter(base_url, logger)
    if not last_chapter:
        logger.error("Không tìm thấy chương cuối, dừng crawl")
        return []
    
    chapters = []
    chapter_number = 1
    consecutive_empty = 0
    max_consecutive_empty = 3
    
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
        while chapter_number <= last_chapter:
            future_to_chapter = {
                executor.submit(fetch_chapter, chap, base_url): chap 
                for chap in range(chapter_number, min(chapter_number + batch_size, last_chapter + 1))
            }
            batch_results = {}
            for future in concurrent.futures.as_completed(future_to_chapter):
                chap_num = future_to_chapter[future]
                try:
                    result = future.result()
                    content = result[1]
                    batch_results[chap_num] = content
                    
                    if content and content not in ["Chương trống", "Không tìm thấy nội dung", "Lỗi 503", "Lỗi 404", "Lỗi không xác định"]:
                        content = clean_text(content)
                        batch_results[chap_num] = content
                        consecutive_empty = 0
                        logger.info(f"Đã crawl thành công chương {chap_num}")
                    else:
                        consecutive_empty += 1
                        logger.warning(f"Chương {chap_num}: {content}")
                        
                except Exception as e:
                    logger.error(f"Lỗi với chương {chap_num}: {str(e)}")
                    batch_results[chap_num] = f"Lỗi: {str(e)}"
                    consecutive_empty += 1
            
            for chap in range(chapter_number, min(chapter_number + batch_size, last_chapter + 1)):
                content = batch_results.get(chap)
                if content:
                    chapters.append((chap, content))
                
                if consecutive_empty >= max_consecutive_empty:
                    logger.warning(f"Đã gặp {max_consecutive_empty} chương trống liên tiếp. Dừng crawl.")
                    return chapters
                    
            chapter_number += batch_size
            time.sleep(2)
            
    # Trước khi return chapters
    missing_chapters = set(range(1, last_chapter + 1)) - set(chap for chap, _ in chapters)
    logger.info(f"Các chương bị bỏ qua: {sorted(missing_chapters)}")
    
    logger.info(f"Hoàn thành crawl {len(chapters)} chương")
    return chapters

def create_story_directory(story_name):
    """Tạo thư mục cho truyện"""
    current_dir = Path(__file__).parent
    story_dir = current_dir / story_name
    story_dir.mkdir(exist_ok=True)
    return story_dir

def save_text_files(chapters, story_dir):
    """Lưu các file text theo chương"""
    current_file = 1
    current_word_count = 0
    current_chapters = []
    target_words = 18000  # Số từ mục tiêu mỗi file
    
    def save_current_file():
        nonlocal current_file, current_word_count, current_chapters
        if not current_chapters:
            return
            
        filename = f"{story_dir.name}_part_{current_file:03d}.txt"
        filepath = story_dir / filename
        
        with open(filepath, 'w', encoding='utf-8') as f:
            for chap_num, content in current_chapters:
                f.write(f"Chương {chap_num}\n")
                f.write("=" * 50 + "\n")
                f.write(content + "\n\n")
        
        print(f"Đã lưu file text: {filename}")
        current_file += 1
        current_word_count = 0
        current_chapters = []
    
    for chap_num, content in chapters:
        # Đếm số từ trong chương
        word_count = len(content.split())
        
        # Nếu thêm chương này vượt quá giới hạn, lưu file hiện tại
        if current_word_count + word_count > target_words and current_chapters:
            save_current_file()
        
        current_chapters.append((chap_num, content))
        current_word_count += word_count
    
    # Lưu file cuối cùng nếu còn chương
    if current_chapters:
        save_current_file()

def create_word_doc(chapters, story_dir):
    """Lưu file Word trong thư mục truyện"""
    try:
        doc = Document()
        for chapter_num, content in chapters:
            doc.add_heading(f"Chương {chapter_num}", level=1)
            doc.add_paragraph(content)
            doc.add_page_break()
        
        output_path = story_dir / f"{story_name}.docx"
        doc.save(str(output_path))
        print(f"Đã lưu file Word tại: {output_path}")
    except Exception as e:
        print(f"Lỗi khi lưu file Word: {e}")
        raise

# Thêm hàm tắt cảnh báo SSL
def disable_ssl_warnings():
    import urllib3
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

if __name__ == "__main__":
    import sys
    
    # Tắt cảnh báo SSL
    disable_ssl_warnings()
    
    if len(sys.argv) > 1:
        base_url = sys.argv[1]
    else:
        print("Vui long nhap URL truyen!")
        print("Vi du: python main.py https://truyenfull.vision/ten-truyen/")
        sys.exit(1)
        
    if not base_url.endswith('/'):
        base_url += '/'
        
    print(f"Bat dau crawl: {base_url}")
    chapters = crawl_story(base_url, batch_size=10, max_workers=5)
    
    if chapters:
        # Lấy tên truyện từ URL
        story_name = base_url.split('/')[-2].replace('-', '_')
        
        # Tạo thư mục cho truyện
        story_dir = create_story_directory(story_name)
        
        try:
            # Lưu file Word
            create_word_doc(chapters, story_dir)
            
            # Lưu các file text
            save_text_files(chapters, story_dir)
            
        except Exception as e:
            print(f"Khong the tao file: {e}")
            sys.exit(1)
    else:
        print("Khong tim thay chuong nao hoac xay ra loi.")
        sys.exit(1)
